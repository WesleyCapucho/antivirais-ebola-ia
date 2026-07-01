import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Estilos de título
    title = doc.add_heading('Protocolo PROSPERO - Meta-Análise', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('TÍTULO DA REVISÃO', style='Heading 1')
    doc.add_paragraph('L-Polymerase Inhibitors Against Filoviruses: A Systematic Review Using Stratified Evidence Tiers')
    
    doc.add_paragraph('OBJETIVO', style='Heading 1')
    doc.add_paragraph('Avaliar a eficácia in vitro e o potencial pan-viral de inibidores da RNA-polimerase dependente de RNA (Proteína L) contra o vírus Ebola (EBOV), vírus Sudão (SUDV) e vírus Marburg (MARV), utilizando uma abordagem estratificada de confiança de evidências.')
    
    doc.add_paragraph('ESTRATÉGIA DE BUSCA (PICO)', style='Heading 1')
    p = doc.add_paragraph()
    p.add_run('Population: ').bold = True
    p.add_run('Estudos in vitro sobre replicação de filovírus.\n')
    p.add_run('Intervention: ').bold = True
    p.add_run('Inibidores da polimerase L (excluindo inibidores primários de VP40/VP35).\n')
    p.add_run('Comparison: ').bold = True
    p.add_run('EBOV-L vs SUDV-L vs MARV-L (validação cross-species).\n')
    p.add_run('Outcomes: ').bold = True
    p.add_run('EC50/IC50 pan-viral, Selectivity Index (SI).')

    doc.add_paragraph('ESTRATIFICAÇÃO DE EVIDÊNCIAS (Inovação Metodológica)', style='Heading 1')
    doc.add_paragraph('Tier 1 (Ouro): Inibição enzimática direta comprovada em polimerase purificada (Confidence Score = 1.0).', style='List Bullet')
    doc.add_paragraph('Tier 2 (Prata): Atividade fenotípica antiviral combinada com validação in silico via Reverse Docking, provando seletividade contra polimerase L vs VP40 (Confidence Score = 0.7).', style='List Bullet')
    doc.add_paragraph('Tier 3 (Bronze): Analogia de classe química (nucleotídeos) sem mecanismo detalhado no estudo (Confidence Score = 0.5).', style='List Bullet')
    
    doc.add_paragraph('CRITÉRIOS DE INCLUSÃO/EXCLUSÃO', style='Heading 1')
    doc.add_paragraph('Inclusão: Artigos revisados por pares (2015-2024), dados quantitativos de EC50/IC50, foco direto ou validado na polimerase L.', style='List Bullet')
    doc.add_paragraph('Exclusão: Ensaios sem dados quantitativos, inibidores de imunologia (mAbs), alvos prioritários documentados como VP40/GP/NP.', style='List Bullet')

    doc.add_paragraph('ANÁLISE DE DADOS E MACHINE LEARNING', style='Heading 1')
    doc.add_paragraph('A meta-análise servirá como dataset curado para treinamento de modelos de Machine Learning. O Confidence Score de cada Tier será diretamente aplicado como "sample_weight" nas Loss Functions (Confidence-Weighted MSE) para garantir aprendizado preferencial sobre dados enzimáticos puros (Tier 1).')

    file_path = os.path.join(r"G:\Meu Drive\PROJETO-ANTIVIRAIS-NANOTECNOLOGIA-IA\1-META-ANALISE\1a-PROTOCOLOS", "protocolo_PROSPERO_refined.docx")
    doc.save(file_path)
    print(f"Protocolo PROSPERO gerado com sucesso em: {file_path}")

if __name__ == "__main__":
    main()
